"""
Create a Word document with embedded AI chat agent.
The document contains a button that launches the chat HTML via VBA macro.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

# Model configuration
MODEL = "ServiceNow-AI/Apriel-1.5-15b-Thinker"

# Self-contained HTML chat app with tool calling
CHAT_HTML = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>PDFClaw AI Chat</title>
<style>
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; background: #0f0f1a; color: #e0e0e0; height: 100vh; display: flex; flex-direction: column; }}
.header {{ background: #1a1a2e; padding: 12px 16px; border-bottom: 1px solid #2a2a4e; }}
.header h1 {{ font-size: 18px; color: #fff; }}
.header p {{ font-size: 10px; color: #888; margin-top: 2px; }}
.api-key-area {{ background: #16213e; padding: 12px 16px; border-bottom: 1px solid #2a2a4e; display: flex; gap: 8px; align-items: center; }}
.api-key-area input {{ flex: 1; background: #0f0f1a; border: 1px solid #2a2a4e; color: #fff; padding: 8px 12px; border-radius: 4px; font-size: 12px; outline: none; }}
.api-key-area input:focus {{ border-color: #e94560; }}
.api-key-area button {{ background: #e94560; color: #fff; border: none; padding: 8px 16px; border-radius: 4px; font-size: 12px; cursor: pointer; }}
.api-key-area button:hover {{ background: #d63851; }}
.api-key-area .saved {{ color: #7fff7f; font-size: 11px; }}
.chat {{ flex: 1; overflow-y: auto; padding: 12px 16px; display: flex; flex-direction: column; gap: 8px; }}
.msg {{ max-width: 85%; padding: 10px 14px; border-radius: 10px; line-height: 1.4; font-size: 13px; white-space: pre-wrap; word-wrap: break-word; }}
.msg.user {{ background: #e94560; color: #fff; align-self: flex-end; border-bottom-right-radius: 4px; }}
.msg.ai {{ background: #16213e; color: #e0e0e0; align-self: flex-start; border-bottom-left-radius: 4px; }}
.msg.system {{ background: #2a2a4e; color: #888; align-self: center; font-size: 11px; text-align: center; }}
.msg.tool {{ background: #1a3a1a; color: #7fff7f; align-self: flex-start; font-size: 11px; font-family: monospace; }}
.msg.error {{ background: #4a1a1a; color: #ff6b6b; }}
.input-area {{ background: #1a1a2e; padding: 12px 16px; border-top: 1px solid #2a2a4e; display: flex; gap: 8px; }}
.input-area input {{ flex: 1; background: #0f0f1a; border: 1px solid #2a2a4e; color: #fff; padding: 10px 14px; border-radius: 6px; font-size: 13px; outline: none; }}
.input-area input:focus {{ border-color: #e94560; }}
.input-area button {{ background: #e94560; color: #fff; border: none; padding: 10px 20px; border-radius: 6px; font-size: 13px; cursor: pointer; font-weight: 600; }}
.input-area button:hover {{ background: #d63851; }}
.input-area button:disabled {{ background: #555; cursor: not-allowed; }}
.typing {{ color: #888; font-style: italic; }}
</style>
</head>
<body>
<div class="header">
    <h1>PDFClaw AI Agent</h1>
    <p>Model: Apriel-1.5-15b-Thinker | Tool Calling | Session Storage</p>
</div>
<div class="api-key-area">
    <input type="password" id="apiKeyInput" placeholder="Enter your Together API key..." />
    <button onclick="saveApiKey()">Save</button>
    <span class="saved" id="apiKeyStatus"></span>
</div>
<div class="chat" id="chat">
    <div class="msg system">Welcome! Enter your Together API key above. Tools: weather, Wikipedia search, fetch URL, calculate, save/load data, time, tweets, generate content.</div>
</div>
<div class="input-area">
    <input type="text" id="input" placeholder="Type your message..." autofocus />
    <button id="sendBtn" onclick="sendMessage()">Send</button>
</div>
<script>
const API_URL = "https://api.together.xyz/v1/chat/completions";
const MODEL = "{MODEL}";

const chatEl = document.getElementById("chat");
const inputEl = document.getElementById("input");
const sendBtn = document.getElementById("sendBtn");
const apiKeyInput = document.getElementById("apiKeyInput");
const apiKeyStatus = document.getElementById("apiKeyStatus");

// Load saved API key
let API_KEY = sessionStorage.getItem('together_api_key') || '';
if (API_KEY) {{
    apiKeyInput.value = API_KEY;
    apiKeyStatus.textContent = 'Saved';
}}

function saveApiKey() {{
    API_KEY = apiKeyInput.value.trim();
    if (API_KEY) {{
        sessionStorage.setItem('together_api_key', API_KEY);
        apiKeyStatus.textContent = 'Saved';
    }} else {{
        apiKeyStatus.textContent = 'Empty';
    }}
}}

// Session storage for user data
const sessionData = JSON.parse(sessionStorage.getItem('pdfclaw_data') || '{{}}');

// Define tools available to the AI
const tools = [
    {{
        type: "function",
        function: {{
            name: "get_weather",
            description: "Get current weather for a city. Use this for weather queries.",
            parameters: {{
                type: "object",
                properties: {{
                    city: {{ type: "string", description: "City name, e.g. 'Nicosia', 'London'" }}
                }},
                required: ["city"]
            }}
        }}
    }},
    {{
        type: "function",
        function: {{
            name: "web_search",
            description: "Search Wikipedia for information. Returns article summary.",
            parameters: {{
                type: "object",
                properties: {{
                    query: {{ type: "string", description: "Search query for Wikipedia" }}
                }},
                required: ["query"]
            }}
        }}
    }},
    {{
        type: "function",
        function: {{
            name: "fetch_url",
            description: "Fetch content from a specific URL. Returns extracted text.",
            parameters: {{
                type: "object",
                properties: {{
                    url: {{ type: "string", description: "Full URL to fetch, e.g. 'https://example.com'" }}
                }},
                required: ["url"]
            }}
        }}
    }},
    {{
        type: "function",
        function: {{
            name: "calculate",
            description: "Perform mathematical calculations",
            parameters: {{
                type: "object",
                properties: {{
                    expression: {{ type: "string", description: "Math expression to evaluate, e.g. '15*200/100'" }}
                }},
                required: ["expression"]
            }}
        }}
    }},
    {{
        type: "function",
        function: {{
            name: "save_data",
            description: "Save data to session storage for later use",
            parameters: {{
                type: "object",
                properties: {{
                    key: {{ type: "string", description: "Name/key for the data" }},
                    value: {{ type: "string", description: "Value to store" }}
                }},
                required: ["key", "value"]
            }}
        }}
    }},
    {{
        type: "function",
        function: {{
            name: "get_data",
            description: "Retrieve previously saved data from session storage",
            parameters: {{
                type: "object",
                properties: {{
                    key: {{ type: "string", description: "Name/key of the data to retrieve" }}
                }},
                required: ["key"]
            }}
        }}
    }},
    {{
        type: "function",
        function: {{
            name: "list_data",
            description: "List all saved data keys in session storage",
            parameters: {{ type: "object", properties: {{}}, required: [] }}
        }}
    }},
    {{
        type: "function",
        function: {{
            name: "get_current_time",
            description: "Get the current date and time",
            parameters: {{ type: "object", properties: {{}}, required: [] }}
        }}
    }},
    {{
        type: "function",
        function: {{
            name: "generate_tweet",
            description: "Generate a tweet-ready message (max 280 chars)",
            parameters: {{
                type: "object",
                properties: {{
                    topic: {{ type: "string", description: "Topic for the tweet" }},
                    style: {{ type: "string", description: "Style: funny, professional, casual, inspirational" }}
                }},
                required: ["topic"]
            }}
        }}
    }},
    {{
        type: "function",
        function: {{
            name: "generate_content",
            description: "Generate various types of content: articles, emails, code, stories, summaries, etc.",
            parameters: {{
                type: "object",
                properties: {{
                    type: {{ type: "string", description: "Type of content: article, email, code, story, summary, list, poem" }},
                    topic: {{ type: "string", description: "Topic or subject of the content" }},
                    length: {{ type: "string", description: "Length: short, medium, long" }},
                    style: {{ type: "string", description: "Style: formal, casual, professional, creative" }}
                }},
                required: ["type", "topic"]
            }}
        }}
    }}
];

// Tool implementations
async function executeTool(toolName, args) {{
    switch(toolName) {{
        case "get_weather":
            try {{
                const geoUrl = 'https://geocoding-api.open-meteo.com/v1/search?name=' + encodeURIComponent(args.city) + '&count=1';
                const geoRes = await fetch(geoUrl);
                if (!geoRes.ok) return 'Error finding city: ' + geoRes.status;
                const geoData = await geoRes.json();
                
                if (!geoData.results || geoData.results.length === 0) {{
                    return 'City "' + args.city + '" not found';
                }}
                
                const lat = geoData.results[0].latitude;
                const lon = geoData.results[0].longitude;
                const cityName = geoData.results[0].name;
                
                const weatherUrl = 'https://api.open-meteo.com/v1/forecast?latitude=' + lat + '&longitude=' + lon + '&current_weather=true';
                const weatherRes = await fetch(weatherUrl);
                if (!weatherRes.ok) return 'Error fetching weather: ' + weatherRes.status;
                const weatherData = await weatherRes.json();
                
                const cw = weatherData.current_weather;
                const temp = cw.temperature;
                const wind = cw.windspeed;
                const desc = cw.weathercode <= 3 ? "Clear" : cw.weathercode <= 49 ? "Cloudy/Foggy" : cw.weathercode <= 99 ? "Rain/Drizzle" : "Storm/Snow";
                
                return 'Weather in ' + cityName + ': ' + temp + 'C, ' + desc + ', Wind: ' + wind + ' km/h';
            }} catch(e) {{
                return 'Error: ' + e.message;
            }}
        
        case "web_search":
            try {{
                const searchUrl = 'https://en.wikipedia.org/api/rest_v1/page/summary/' + encodeURIComponent(args.query);
                const res = await fetch(searchUrl);
                if (!res.ok) {{
                    const searchApiUrl = 'https://en.wikipedia.org/w/api.php?action=opensearch&search=' + encodeURIComponent(args.query) + '&limit=3&format=json&origin=*';
                    const searchRes = await fetch(searchApiUrl);
                    if (!searchRes.ok) return 'Wikipedia search failed: ' + searchRes.status;
                    const searchData = await searchRes.json();
                    if (searchData[1] && searchData[1].length > 0) {{
                        let results = 'Wikipedia results for "' + args.query + '":\\n';
                        for (let i = 0; i < searchData[1].length; i++) {{
                            results += (i+1) + '. ' + searchData[1][i] + ' - ' + searchData[3][i] + '\\n';
                        }}
                        return results;
                    }}
                    return 'No Wikipedia results for "' + args.query + '"';
                }}
                const data = await res.json();
                return 'Wikipedia: ' + data.title + '\\n' + (data.extract || data.description || 'No summary available') + '\\nSource: ' + (data.content_urls?.desktop?.page || 'Wikipedia');
            }} catch(e) {{
                return 'Error: ' + e.message;
            }}
        
        case "fetch_url":
            try {{
                let url = args.url;
                if (!url.startsWith('http')) url = 'https://' + url;
                
                const corsProxies = [
                    'https://api.allorigins.win/raw?url=' + encodeURIComponent(url),
                    'https://corsproxy.io/?' + encodeURIComponent(url)
                ];
                
                let html = null;
                for (const proxyUrl of corsProxies) {{
                    try {{
                        const res = await fetch(proxyUrl);
                        if (res.ok) {{
                            html = await res.text();
                            break;
                        }}
                    }} catch(e) {{}}
                }}
                
                if (!html) return 'Could not fetch ' + url + '. CORS proxies blocked. Try Wikipedia search instead.';
                
                const text = html.replace(/<script[^>]*>[\\s\\S]*?<\\/script>/gi, '')
                               .replace(/<style[^>]*>[\\s\\S]*?<\\/style>/gi, '')
                               .replace(/<[^>]+>/g, ' ')
                               .replace(/&nbsp;/g, ' ')
                               .replace(/&amp;/g, '&')
                               .replace(/\\s+/g, ' ')
                               .substring(0, 3000);
                return 'Content from ' + url + ':\\n' + text;
            }} catch(e) {{
                return 'Error: ' + e.message;
            }}
        
        case "calculate":
            try {{
                const result = Function('"use strict"; return (' + args.expression + ')')();
                return 'Result: ' + args.expression + ' = ' + result;
            }} catch(e) {{
                return 'Error calculating: ' + e.message;
            }}
        
        case "save_data":
            sessionData[args.key] = args.value;
            sessionStorage.setItem('pdfclaw_data', JSON.stringify(sessionData));
            return 'Saved: ' + args.key + ' = "' + args.value + '"';
        
        case "get_data":
            const value = sessionData[args.key];
            if (value !== undefined) {{
                return 'Retrieved: ' + args.key + ' = "' + value + '"';
            }}
            return 'No data found for key: ' + args.key;
        
        case "list_data":
            const keys = Object.keys(sessionData);
            if (keys.length === 0) return "No saved data found.";
            return 'Saved data keys: ' + keys.join(', ') + '\\nData: ' + JSON.stringify(sessionData, null, 2);
        
        case "get_current_time":
            const now = new Date();
            return 'Current time: ' + now.toLocaleString();
        
        case "generate_tweet":
            return 'Generate a tweet about "' + args.topic + '" (style: ' + (args.style || 'casual') + '). Keep it under 280 characters. Return ONLY the tweet text.';
        
        case "generate_content":
            const length = args.length || 'medium';
            const style = args.style || 'professional';
            return 'Generate a ' + length + ' ' + args.type + ' about "' + args.topic + '" in ' + style + ' style. Return ONLY the content, no explanations.';
        
        default:
            return 'Unknown tool: ' + toolName;
    }}
}}

let messages = [
    {{ 
        role: "system", 
        content: "You are a helpful AI assistant with access to tools. When you need to search the web, calculate, save/retrieve data, or perform specific tasks, use the appropriate tool. Always be helpful and concise. After using a tool, explain the result to the user naturally."
    }}
];

inputEl.addEventListener("keydown", function(e) {{
    if (e.key === "Enter" && !sendBtn.disabled) sendMessage();
}});

function addMsg(text, cls) {{
    const div = document.createElement("div");
    div.className = "msg " + cls;
    div.textContent = text;
    chatEl.appendChild(div);
    chatEl.scrollTop = chatEl.scrollHeight;
    return div;
}}

async function sendMessage() {{
    const msg = inputEl.value.trim();
    if (!msg) return;
    
    if (!API_KEY) {{
        addMsg("Please enter your Together API key above first!", "system");
        apiKeyInput.focus();
        return;
    }}
    
    addMsg(msg, "user");
    inputEl.value = "";
    sendBtn.disabled = true;
    
    messages.push({{ role: "user", content: msg }});
    
    const typing = addMsg("Thinking...", "ai typing");
    
    try {{
        let response = await callAI();
        typing.textContent = response;
        typing.className = "msg ai";
    }} catch(e) {{
        typing.textContent = "Error: " + e.message;
        typing.className = "msg ai error";
    }}
    
    sendBtn.disabled = false;
    inputEl.focus();
}}

async function callAI() {{
    const res = await fetch(API_URL, {{
        method: "POST",
        headers: {{
            "Content-Type": "application/json",
            "Authorization": "Bearer " + API_KEY
        }},
        body: JSON.stringify({{
            model: MODEL,
            messages: messages,
            tools: tools,
            tool_choice: "auto",
            max_tokens: 1024
        }})
    }});
    
    if (!res.ok) {{
        const errText = await res.text();
        throw new Error('API Error: ' + res.status + ' - ' + errText.substring(0, 100));
    }}
    
    const data = await res.json();
    const assistantMsg = data.choices[0].message;
    
    messages.push(assistantMsg);
    
    if (assistantMsg.tool_calls && assistantMsg.tool_calls.length > 0) {{
        for (const toolCall of assistantMsg.tool_calls) {{
            const toolName = toolCall.function.name;
            const args = JSON.parse(toolCall.function.arguments);
            addMsg('Tool: ' + toolName + '(' + JSON.stringify(args) + ')', "tool");
            
            const result = await executeTool(toolName, args);
            addMsg('Result: ' + result, "tool");
            
            messages.push({{
                role: "tool",
                tool_call_id: toolCall.id,
                content: result
            }});
        }}
        
        const finalRes = await fetch(API_URL, {{
            method: "POST",
            headers: {{
                "Content-Type": "application/json",
                "Authorization": "Bearer " + API_KEY
            }},
            body: JSON.stringify({{
                model: MODEL,
                messages: messages,
                max_tokens: 1024
            }})
        }});
        
        if (!finalRes.ok) throw new Error("Failed to get final response");
        
        const finalData = await finalRes.json();
        const finalMsg = finalData.choices[0].message;
        messages.push(finalMsg);
        return finalMsg.content;
    }}
    
    return assistantMsg.content || "No response";
}}
</script>
</body>
</html>"""

# VBA Macro code to embed in Word
VBA_MACRO = '''
Sub LaunchAIChat()
    ' Extract embedded HTML and open in browser
    Dim htmlPath As String
    Dim fso As Object
    Dim file As Object
    Dim htmlContent As String
    
    ' Create temp file
    htmlPath = Environ("TEMP") & "\\pdfclaw_chat.html"
    
    ' HTML content (embedded)
    htmlContent = "EMBEDDED_HTML_PLACEHOLDER"
    
    ' Write to file
    Set fso = CreateObject("Scripting.FileSystemObject")
    Set file = fso.CreateTextFile(htmlPath, True)
    file.Write htmlContent
    file.Close
    
    ' Open in default browser
    CreateObject("Shell.Application").Open htmlPath
End Sub
'''

def create_word_document():
    """Create Word document with instructions and button placeholder"""
    doc = Document()
    
    # Title
    title = doc.add_heading('PDFClaw AI Agent for Word', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('AI Chat Agent Embedded in Word Document')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    
    # Model info
    model_para = doc.add_paragraph()
    model_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    model_run = model_para.add_run('Model: Apriel-1.5-15b-Thinker | Tool Calling Enabled')
    model_run.font.size = Pt(12)
    model_run.font.color.rgb = RGBColor(0xe9, 0x45, 0x60)
    model_run.bold = True
    
    doc.add_paragraph()
    
    # Instructions box
    doc.add_heading('How to Use:', level=1)
    
    instructions = [
        '1. Enable macros in Word (File > Options > Trust Center > Macro Settings > Enable all macros)',
        '2. Press the "Launch AI Chat" button below',
        '3. Your browser will open with the AI chat interface',
        '4. Enter your Together API key (get one free at api.together.xyz)',
        '5. Start chatting with the AI!'
    ]
    
    for inst in instructions:
        p = doc.add_paragraph(inst)
        p.paragraph_format.left_indent = Inches(0.5)
    
    doc.add_paragraph()
    
    # Button placeholder
    button_para = doc.add_paragraph()
    button_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    button_run = button_para.add_run('[LAUNCH AI CHAT BUTTON]')
    button_run.font.size = Pt(16)
    button_run.font.color.rgb = RGBColor(0xe9, 0x45, 0x60)
    button_run.bold = True
    
    doc.add_paragraph()
    button_note = doc.add_paragraph('Note: Add a button shape and assign the "LaunchAIChat" macro to it.')
    button_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    button_note.runs[0].font.italic = True
    button_note.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    
    doc.add_paragraph()
    
    # Features section
    doc.add_heading('Available Tools:', level=1)
    
    tools = [
        ('get_weather', 'Get current weather for any city'),
        ('web_search', 'Search Wikipedia for information'),
        ('fetch_url', 'Fetch content from any URL'),
        ('calculate', 'Perform mathematical calculations'),
        ('save_data / get_data', 'Store and retrieve data in session'),
        ('get_current_time', 'Get current date and time'),
        ('generate_tweet', 'Generate tweet-ready content'),
        ('generate_content', 'Generate articles, emails, code, stories'),
    ]
    
    for tool, desc in tools:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        tool_run = p.add_run(f'{tool}: ')
        tool_run.bold = True
        p.add_run(desc)
    
    doc.add_paragraph()
    
    # Example commands
    doc.add_heading('Example Commands:', level=1)
    
    examples = [
        '"What is the weather in Nicosia?"',
        '"Search Wikipedia for Cyprus"',
        '"Fetch https://example.com"',
        '"Calculate 15% of 250"',
        '"Save my name as John"',
        '"Generate an article about AI"',
    ]
    
    for ex in examples:
        p = doc.add_paragraph(ex)
        p.paragraph_format.left_indent = Inches(0.5)
        p.runs[0].font.italic = True
    
    # Save document
    doc.save('pdfclaw_word_agent.docx')
    print("Word document created: pdfclaw_word_agent.docx")

def save_html_file():
    """Save the HTML chat app as a separate file"""
    with open('chat_app.html', 'w', encoding='utf-8') as f:
        f.write(CHAT_HTML)
    print("HTML chat app saved: chat_app.html")

def save_vba_module():
    """Save VBA module with embedded HTML"""
    # Escape the HTML for VBA string
    html_escaped = CHAT_HTML.replace('"', '""').replace('\n', '" & vbCrLf & "')
    
    vba_code = VBA_MACRO.replace('EMBEDDED_HTML_PLACEHOLDER', html_escaped)
    
    with open('ai_chat_macro.bas', 'w', encoding='utf-8') as f:
        f.write("Attribute VB_Name = \"AIChatModule\"\n")
        f.write(vba_code)
    print("VBA module saved: ai_chat_macro.bas")

if __name__ == "__main__":
    create_word_document()
    save_html_file()
    save_vba_module()
    print("\nFiles created in word/ directory:")
    print("  - pdfclaw_word_agent.docx (Word document)")
    print("  - chat_app.html (Standalone HTML chat app)")
    print("  - ai_chat_macro.bas (VBA module to import)")
